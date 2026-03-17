"""
Script para gerar o documento Word do Desafio 2.
Execute: python gerar_desafio2.py
"""
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ──────────────────────────────────────────────────────────────────
def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    row.cells[0].text = col1
    row.cells[1].text = col2
    if header:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

# ── carregar e processar dados ────────────────────────────────────────────────
print("Carregando dados...")
df = pd.read_csv(
    'assets/vendas_linha_petshop_2020_2024.csv',
    sep=';', encoding='latin-1', low_memory=False
)

for col in ['valor', 'quantidade', 'valor_total_bruto', 'valor_comissao', 'lucro_liquido']:
    df[col] = df[col].astype(str).str.replace(',', '.').str.strip()
    df[col] = pd.to_numeric(df[col], errors='coerce')

# Outliers por IQR
numeric_cols = ['valor', 'quantidade', 'valor_total_bruto', 'valor_comissao', 'lucro_liquido']
total_outliers_mask = pd.Series([False] * len(df))
outlier_details = {}
for col in numeric_cols:
    Q1 = df[col].quantile(0.25)
    Q3 = df[col].quantile(0.75)
    IQR = Q3 - Q1
    lower = Q1 - 1.5 * IQR
    upper = Q3 + 1.5 * IQR
    mask = (df[col] < lower) | (df[col] > upper)
    total_outliers_mask = total_outliers_mask | mask
    outlier_details[col] = {
        'Q1': Q1, 'Q3': Q3, 'IQR': IQR,
        'lower': lower, 'upper': upper, 'count': int(mask.sum())
    }

total_outliers = int(total_outliers_mask.sum())
pct_outliers = total_outliers / len(df) * 100
total_rows = len(df)

print(f"Total linhas: {total_rows}")
print(f"Total outliers: {total_outliers} ({pct_outliers:.1f}%)")
print("Criando documento Word...")

# ── criar documento ───────────────────────────────────────────────────────────
doc = Document()

# titulo
title = doc.add_heading('Desafio 2 – Explorando e Analisando Dados para uso em Ciência de Dados', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Empresa: Melhores Compras | Dataset: vendas_linha_petshop_2020_2024.csv')
doc.add_paragraph()

# ─── PERGUNTA 1 ───────────────────────────────────────────────────────────────
set_heading(doc, 'Pergunta 1 – Estratégia para criação da base de dados de treinamento, teste e validação', level=1)

add_paragraph(doc,
    'Para construir um modelo preditivo robusto, os dados disponíveis devem ser divididos em '
    'três conjuntos mutuamente exclusivos: treinamento, teste e validação. A estratégia '
    'recomendada para o dataset da Melhores Compras é a seguinte:')

# tabela de proporções
t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
hdr = t.rows[0].cells
hdr[0].text = 'Conjunto'
hdr[1].text = 'Proporção'
hdr[2].text = 'Finalidade'
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True

dados = [
    ('Treinamento', '70 %',
     'Utilizado para o ajuste dos parâmetros do modelo (coeficientes da regressão). '
     'Quanto maior essa fatia, mais o modelo aprende os padrões dos dados.'),
    ('Teste', '20 %',
     'Avalia o desempenho do modelo em dados não vistos durante o treinamento. '
     'Fornece a estimativa honesta de generalização (R², RMSE).'),
    ('Validação', '10 %',
     'Reservado para ajuste de hiperparâmetros e seleção de modelo, sem contaminar '
     'a avaliação final realizada na base de teste.'),
]
for nome, prop, desc in dados:
    row = t.add_row().cells
    row[0].text = nome
    row[1].text = prop
    row[2].text = desc

doc.add_paragraph()
add_paragraph(doc,
    'Justificativa: A proporção 70/20/10 é amplamente adotada na literatura de Machine Learning '
    'e também referenciada nos materiais do curso (Cap. 10 – Modelagem Preditiva). '
    'Com 250.059 registros disponíveis, cada conjunto terá, respectivamente, '
    '~175.041, ~50.012 e ~25.006 linhas, volumes suficientes para treinar e avaliar '
    'modelos de forma estatisticamente confiável. O parâmetro random_state=42 garante '
    'reprodutibilidade na separação aleatória, conforme boas práticas do sklearn.')

# ─── PERGUNTA 2 ───────────────────────────────────────────────────────────────
set_heading(doc, 'Pergunta 2 – Variáveis numéricas e categóricas', level=1)

add_paragraph(doc,
    'Variável numérica (quantitativa): representa quantidades mensuráveis, permite '
    'operações aritméticas (soma, média) e pode ser contínua (ex.: valor monetário) '
    'ou discreta (ex.: contagem de itens). Variável categórica (qualitativa): representa '
    'grupos ou categorias sem escala numérica natural. Podem ser nominais (sem ordem, '
    'ex.: forma de pagamento) ou ordinais (com ordem, ex.: nível de risco).')

doc.add_paragraph()
add_paragraph(doc, 'Classificação das variáveis do arquivo vendas_linha_petshop_2020_2024.csv:', bold=True)

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
hdr2 = t2.rows[0].cells
hdr2[0].text = 'Coluna'
hdr2[1].text = 'Tipo'
hdr2[2].text = 'Justificativa'
for cell in hdr2:
    for run in cell.paragraphs[0].runs:
        run.bold = True

colunas = [
    ('cod_pedido', 'Numérica (discreta) *', 'Identificador inteiro único de pedido. *Apesar de numérico, não deve ser usado como feature preditiva.'),
    ('valor', 'Numérica (contínua)', 'Valor unitário do produto em R$ – escala contínua de preços.'),
    ('quantidade', 'Numérica (discreta)', 'Contagem de unidades vendidas.'),
    ('valor_total_bruto', 'Numérica (contínua)', 'Multiplicação de valor × quantidade – escala monetária contínua.'),
    ('valor_comissao', 'Numérica (contínua)', 'Valor da comissão em R$ – escala contínua.'),
    ('lucro_liquido', 'Numérica (contínua)', 'Lucro líquido em R$ após impostos – escala contínua.'),
    ('regiao_pais', 'Categórica (nominal)', '5 regiões sem ordem natural: Norte, Sul, Nordeste, Sudeste, Centro-Oeste.'),
    ('produto', 'Categórica (nominal)', 'Nome do produto – mais de 50.000 SKUs distintos.'),
    ('data', 'Data/Hora', 'Timestamp do pedido. Pode ser decomposta em dia, mês, ano, dia da semana (numérico ou categórico).'),
    ('estado', 'Categórica (nominal)', '25 estados brasileiros sem ordem natural.'),
    ('formapagto', 'Categórica (nominal)', '5 formas de pagamento: Cartão Crédito, Débito, Pix, Boleto, Dinheiro.'),
    ('centro_distribuicao', 'Categórica (nominal)', '5 centros de distribuição sem hierarquia.'),
    ('responsavelpedido', 'Categórica (nominal)', 'Nome do responsável pelo pedido – alta cardinalidade.'),
    ('categoriaprod', 'Categórica (nominal)', '7 categorias de produto: Alimentação, Brinquedo, Acessório etc.'),
]
for nome, tipo, just in colunas:
    row = t2.add_row().cells
    row[0].text = nome
    row[1].text = tipo
    row[2].text = just

# ─── PERGUNTA 3 ───────────────────────────────────────────────────────────────
set_heading(doc, 'Pergunta 3 – Importância do cod_pedido em análises preditivas', level=1)

add_paragraph(doc,
    'O cod_pedido é um identificador técnico único gerado pelo sistema de e-commerce '
    'para cada pedido transacionado. Sua importância em análises preditivas é, '
    'paradoxalmente, a de ser excluído como variável independente (feature) do modelo, '
    'pelos seguintes motivos:')

bullets = [
    ('Sem significado preditivo:', 'O número sequencial do pedido não carrega '
     'informação de negócio capaz de explicar ou prever valores de vendas, lucro '
     'ou comportamento do consumidor.'),
    ('Risco de overfitting:', 'Incluir um identificador único como feature cria '
     'um mapeamento artificial pedido→resultado que se ajusta perfeitamente ao '
     'treino, mas falha completamente em dados novos.'),
    ('Duplicidade controlada:', f'No dataset foram encontrados {250059:,} registros '
     'com {224918:,} códigos únicos (25.141 duplicatas). As duplicatas ocorrem porque '
     'um mesmo pedido pode conter múltiplos itens/linhas. Isso confirma que o campo '
     'identifica o pedido, não a linha de venda.'),
    ('Uso correto:', 'O cod_pedido deve ser utilizado apenas para operações de '
     'JOIN, deduplicação, rastreabilidade e auditoria dos dados, nunca como feature '
     'de entrada em modelos de Machine Learning.'),
]
for titulo, texto in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(titulo + ' ')
    run_b.bold = True
    p.add_run(texto)

# ─── PERGUNTA 4 ───────────────────────────────────────────────────────────────
set_heading(doc, 'Pergunta 4 – Identificação de outliers', level=1)

add_paragraph(doc,
    'Para identificar os outliers foram utilizados dois critérios complementares, '
    'ambos amplamente discutidos nos materiais do curso (Cap. 10 – Modelagem Preditiva):')

p = doc.add_paragraph(style='List Number')
run = p.add_run('Método IQR (Intervalo Interquartil): ')
run.bold = True
p.add_run(
    'Para cada variável numérica, calculou-se o primeiro quartil (Q1) e o terceiro '
    'quartil (Q3). O Intervalo Interquartil é IQR = Q3 – Q1. Registros com valor '
    'fora do intervalo [Q1 – 1,5×IQR, Q3 + 1,5×IQR] são considerados outliers. '
    'Essa abordagem é robusta a distribuições assimétricas, comuns em dados de vendas.'
)

p2 = doc.add_paragraph(style='List Number')
run2 = p2.add_run('Critério aplicado: ')
run2.bold = True
p2.add_run('Registros que apresentam pelo menos uma variável numérica fora do intervalo IQR.')

doc.add_paragraph()
add_paragraph(doc, 'Resultados por variável:', bold=True)

t3 = doc.add_table(rows=1, cols=6)
t3.style = 'Table Grid'
hdr3 = t3.rows[0].cells
for i, h in enumerate(['Variável', 'Q1', 'Q3', 'IQR', 'Limite inferior', 'Outliers']):
    hdr3[i].text = h
    for run in hdr3[i].paragraphs[0].runs:
        run.bold = True

for col, d in outlier_details.items():
    row = t3.add_row().cells
    row[0].text = col
    row[1].text = f"{d['Q1']:,.2f}"
    row[2].text = f"{d['Q3']:,.2f}"
    row[3].text = f"{d['IQR']:,.2f}"
    row[4].text = f"{d['lower']:,.2f}"
    row[5].text = str(d['count'])

doc.add_paragraph()
add_paragraph(doc,
    f'Total de registros com pelo menos um valor discrepante (outlier): '
    f'{total_outliers:,} registros ({pct_outliers:.1f}% do total de {total_rows:,} linhas). '
    f'O código Python a seguir reproduz a análise:', bold=False)

# bloco de código
p_code = doc.add_paragraph()
p_code.style = 'No Spacing'
run_code = p_code.add_run(
    "Q1 = df[col].quantile(0.25)\n"
    "Q3 = df[col].quantile(0.75)\n"
    "IQR = Q3 - Q1\n"
    "outliers = (df[col] < Q1 - 1.5*IQR) | (df[col] > Q3 + 1.5*IQR)"
)
run_code.font.name = 'Courier New'
run_code.font.size = Pt(9)

# ─── PERGUNTA 5 ───────────────────────────────────────────────────────────────
set_heading(doc, 'Pergunta 5 – Estratégia para tratar os outliers identificados', level=1)

add_paragraph(doc,
    'Existem diversas estratégias para tratamento de outliers. A escolha ideal '
    'depende da natureza do dado e do objetivo do modelo. Para o contexto da '
    'Melhores Compras, propõem-se as seguintes abordagens, em ordem de preferência:')

estrategias = [
    ('1. Investigação e contextualização (passo obrigatório):',
     'Antes de qualquer tratamento, é essencial verificar se os outliers são erros '
     'de entrada de dados (ex.: quantidade = -1.000 ou quantidade = 1.000) ou eventos '
     'reais (ex.: compras em grande volume por revendedores). No dataset, foram '
     'encontradas quantidades negativas (-1.000) e extremamente altas (1.000), '
     'sugerindo possíveis erros de digitação que devem ser investigados.'),
    ('2. Remoção (para erros comprovados):',
     'Registros cujos valores sejam claramente inconsistentes com o domínio do negócio '
     '(ex.: quantidade negativa, valor unitário = R$ 0,00) devem ser removidos após '
     'validação da equipe de negócio. Cuidado: a remoção deve ser documentada e não '
     'deve reduzir o dataset de forma prejudicial ao modelo.'),
    ('3. Winsorização / Capping (para distribuições com caudas longas):',
     'Substituir valores extremos pelos percentis limites (ex.: P1 e P99). '
     'Essa técnica preserva a linha no dataset enquanto suaviza o impacto dos extremos '
     'no treinamento do modelo, sendo recomendada para variáveis como valor_comissao '
     'e lucro_liquido, que apresentam caudas muito longas.'),
    ('4. Transformação logarítmica (para variáveis fortemente assimétricas):',
     'Aplicar log(1 + x) nas variáveis monetárias (valor, valor_total_bruto, '
     'lucro_liquido) comprime a escala e reduz a influência de valores extremos, '
     'melhorando o desempenho de modelos lineares.'),
    ('5. Imputação pela mediana (para outliers em variáveis com missings):',
     'Para variáveis como valor (99.673 valores nulos após conversão), '
     'recomenda-se imputar a mediana por categoria de produto, evitando distorção '
     'causada por outliers que afetariam a média.'),
]

for titulo, texto in estrategias:
    p = doc.add_paragraph()
    run_b = p.add_run(titulo + '\n')
    run_b.bold = True
    p.add_run(texto)

doc.add_paragraph()
add_paragraph(doc,
    'Recomendação final: para o modelo preditivo da Melhores Compras, a estratégia '
    'recomendada é (1) investigar e documentar os outliers com a equipe de negócio, '
    '(2) remover erros evidentes de digitação, e (3) aplicar winsorização nos '
    'percentis 1% e 99% nas variáveis monetárias antes do treinamento do modelo. '
    'Toda transformação deve ser registrada no pipeline de preparação dos dados '
    'para garantir rastreabilidade e reprodutibilidade.', bold=False)

# ── salvar ───────────────────────────────────────────────────────────────────
output_path = 'Desafio2_Analise_Dados.docx'
doc.save(output_path)
print(f"\nDocumento gerado: {output_path}")
